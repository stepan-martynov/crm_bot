from datetime import datetime
from babel.dates import format_date
import os
from pprint import pprint

from docx import Document
from docx.shared import Inches
from docxtpl import DocxTemplate

from PIL import Image
from num2words import num2words

DOCX_TEMPLATE_PATH = f'{os.getcwd()}/E&V2.docx'
base_dir_path = 'D:/YandexDisk/С. Мартынов'

# =========== INPUT ===========
contract_dir = '223 Невский 152'
price = 3000
service_name = 'Профессиональная фотосъемка недвижимости'
location = 'Санкт-Петербург'


# ========== END INPUT ==========
# TODO Запрос к папке на ЯД

def date_format_to_ru(d: datetime) -> str:
    return format_date(d, format='long', locale='ru')[:-3]


litter_price = num2words(price, lang="ru")
contract_number = f'ev-sm-{contract_dir.split(" ")[0]}'

image_dir_path = f'{base_dir_path}/{contract_dir}/на сайт/'
image_list = [f for f in os.listdir(image_dir_path) if '.jpg' in f]

contract_date = date_format_to_ru(datetime.strptime(image_list[0][:8], "%Y%m%d"))  # '9 сентября 2020'
cur_date = date_format_to_ru(datetime.now())
print(image_list)

context = {
    'num': contract_number,
    'NUM': contract_number.upper(),
    'date': contract_date,
    'DATE': contract_date.upper(),
    'cost': str(price),
    'letter_cost': litter_price,
    'service_name': service_name,
    'final_cost': str(price),
    'cur_date': cur_date,
    'adress': f'{location}, {contract_dir[4:]}',
    'photos': '{{photos}}',
    'num_of_photos': str(len(image_list))
}


def add_images(image_list, doc):
    for p in doc.paragraphs:
        if '{{photos}}' in p.text:
            p.text = ''
            run = p.add_run()
            for image in image_list:
                try:
                    img = Image.open(os.path.join(base_dir_path, contract_dir, "на сайт", image))
                    (width, height) = (img.width // 5, img.height // 5)
                    img.resize((width, height)).save("buff.jpg")
                    if width > height:
                        run.add_picture("buff.jpg", width=Inches(1))
                    else:
                        run.add_picture("buff.jpg", height=Inches(0.66))
                except:
                    print(f'{os.path.join(image_dir_path, image)} is not image')
            os.remove('buff.jpg')

    return doc


if __name__ == '__main__':
    doc = DocxTemplate(DOCX_TEMPLATE_PATH)
    pprint(context)

    doc.render(context)

    base_doc_path = f'{base_dir_path}/{contract_dir}/{contract_number}.docx'

    try:
        doc.save(base_doc_path)
    except Exception as e:
        print(e)

    doc = Document(base_doc_path)
    add_images(image_list, doc)
    doc.save(base_doc_path)
