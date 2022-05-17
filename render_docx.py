import os
import datetime
import pprint

from PIL import Image
from loguru import logger
from dataclasses import dataclass, field, asdict
from docx import Document
from docxtpl import DocxTemplate
from docx.shared import Cm
from babel.dates import format_date
from num2words import num2words

import yadiskapi

logger.add('test.log')


def parse_full_name(s):
    pass


def set_doc_name_string(s):
    pass


def parse_expression(s, exp_s):
    pass


@dataclass
class Customer:
    name: str
    full_name: str = field(init=False)
    doc_name_string: str = field(init=False)
    inn: str = field(init=False)
    kpp: str = field(init=False)

    def __post_init__(self):
        object.__setattr__(self, 'full_name', parse_full_name(self.raw_string))
        object.__setattr__(self, 'doc_name_string', set_doc_name_string(self))
        object.__setattr__(self, 'inn', parse_expression(self.raw_string, 'inn_str'))
        object.__setattr__(self, 'kpp', parse_expression(self.raw_string, 'kpp_str'))

    @property
    def raw_string(self):
        s: str = ''
        return s


@dataclass
class Entity:
    url: str
    broker: str
    customer: str
    service: str
    price: str

    @property
    def photo_dir(self) -> yadiskapi.YaDiskObject:
        return yadiskapi.get_public_meta(self.url)

    @property
    def img_list(self) -> list[yadiskapi.YaDiskObject]:
        return yadiskapi.create_img_list(self.url, img_list=[], names_list=[])


def _get_num(photo_dir: yadiskapi.YaDiskObject) -> str:
    return f'ev-sm-{photo_dir.name.split()[0]}'


def _get_litter_price(price: str) -> str:
    return num2words(price, lang="ru")


def date_format_to_ru(d: datetime) -> str:
    return format_date(d, format='long', locale='ru')[:-3]


def _get_address(service: str, address: str) -> str:
    return f'{service}, {address[4:]}'


@dataclass
class OutEntity(Entity):
    num: str = field(init=False)
    upper_num: str = field(init=False)
    date: str = field(init=False, default='')
    upper_date: str = field(init=False)
    letter_price: str = field(init=False)
    cur_date: str = field(init=False)
    address: str = field(init=False)
    photo_counter: str = field(init=False)
    photos: str = '{{photos}}'

    def __post_init__(self):
        object.__setattr__(self, 'num', _get_num(self.photo_dir))
        object.__setattr__(self, 'upper_num', self.num.upper())
        object.__setattr__(self, 'date', date_format_to_ru(self.shooting_date))
        object.__setattr__(self, 'upper_date', self.date.upper())
        object.__setattr__(self, 'letter_price', _get_litter_price(self.price))
        object.__setattr__(self, 'cur_date', date_format_to_ru(datetime.datetime.now()))
        object.__setattr__(self, 'address', _get_address(self.service, self.photo_dir.name))
        object.__setattr__(self, 'photo_counter', str(len(self.img_list)))

    @property
    def shooting_date(self) -> datetime:
        return datetime.datetime.strptime(self.img_list[0].name[:8], "%Y%m%d")

    @property
    def exel_string(self) -> str:
        return f'{self.shooting_date.strftime("%d.%m.%Y")}\t\t{self.photo_dir.name[4:]}\t\t\t{self.broker}\t\t\t{self.num}' \
               f' не оплачено\t{self.price}\t\t{self.url}'

    @property
    def portal_string(self) -> str:
        return f'Ссылка на объект: {self.url}\n' \
               f'Ответственный за фотосессию: {self.broker}\n' \
               f'{self.num}\n' \
               f'Фотосъемка квартиры по адресу {self.address}'


def add_images(image_list, doc):
    for p in doc.paragraphs:
        if '{{photos}}' in p.text:
            p.text = ''
            run = p.add_run()
            for image in image_list:
                f = yadiskapi.download_preview(image)
                try:
                    with Image.open(f) as img:
                        img.save(f, "JPEG")
                        if img.width > img.height:
                            run.add_picture(f, width=Cm(2.7))
                        else:
                            run.add_picture(f, height=Cm(1.8))
                    os.remove(f)
                except Exception as e:
                    logger.error(e)
    return doc


def render_doc(entity: OutEntity):
    doc_template = f'{entity.customer}.docx'
    DOCX_TEMPLATE_PATH = os.path.join(os.getcwd(), 'docs_tamplates', doc_template)
    doc = DocxTemplate(DOCX_TEMPLATE_PATH)
    doc.render(asdict(entity))
    file_name = f'{entity.num}.docx'
    doc.save(file_name)

    doc = Document(file_name)
    add_images(entity.img_list, doc)
    doc.save(file_name)


def main():
    # pass
    #
    ent1 = OutEntity(
        url="https://disk.yandex.ru/d/KLNVNKhQ5REaJw",
        broker="Анна Романова",
        customer="E&V privet",
        service="Город",
        price='2000'
    )

    pprint.pprint(asdict(ent1))
    # render_doc(ent1)


if __name__ == '__main__':
    main()
